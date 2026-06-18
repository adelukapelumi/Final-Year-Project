from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PROJECT_DIR = ROOT / "Final Year Project by Codex"
DRAFT_PATH = PROJECT_DIR / "Draft of Final Year Project with Mendeley cite.docx"
OUTPUT_PATH = PROJECT_DIR / "Defense Ready Final Year Project Report.docx"
BACKUP_DIR = PROJECT_DIR / "report_work" / "backups"
ORDERED_JSON_PATH = PROJECT_DIR / "report_work" / "draft_ordered.json"
DIAGRAMS_DIR = PROJECT_DIR / "report_assets" / "diagrams"
SCREENSHOTS_DIR = PROJECT_DIR / "report_assets" / "screenshots"
BENCHMARK_JSON_PATH = PROJECT_DIR / "benchmark_results.json"


CHAPTER_MAP = {
    "CHAPTER ONE": "CHAPTER ONE: INTRODUCTION",
    "CHAPTER TWO": "CHAPTER TWO: LITERATURE REVIEW",
    "CHAPTER THREE": "CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN",
    "CHAPTER FOUR": "CHAPTER FOUR: SYSTEM IMPLEMENTATION, EVALUATION, AND DISCUSSION",
}

SKIP_PARAGRAPHS = {
    "INTRODUCTION",
    "LITERATURE REVIEW",
    "SYSTEM ANALYSIS AND DESIGN",
    "IMPLEMENTATION, EVALUATION, AND INTERFACES",
    "SYSTEM IMPLEMENTATION, EVALUATION, AND DISCUSSION",
}

OBJECTIVE_TEXTS = {
    "To review existing e-voting systems and the cryptographic approaches they use.",
    "To analyse the security requirements and threat model for diaspora e-voting.",
    "To design a zk-STARK-based ballot validity constraint system.",
    "To implement a working prototype integrating zk-STARK proof generation, encrypted ballot submission, and public verification.",
    "To evaluate the prototype’s performance and security properties.",
}

FIGURE_ENTRIES = [
    ("fig_3_1", "Figure 3.1  Use-case diagram for DiasporaVote actors"),
    ("fig_3_2", "Figure 3.2  DiasporaVote system architecture"),
    ("fig_3_3", "Figure 3.3  Voting workflow for the controlled referendum prototype"),
    ("fig_3_4", "Figure 3.4  zk-STARK proof verification flow"),
    ("fig_3_5", "Figure 3.5  Database and entity-relationship design"),
    ("fig_4_1", "Figure 4.1  Landing page of the DiasporaVote prototype"),
    ("fig_4_2", "Figure 4.2  Mock NIN accreditation page"),
    ("fig_4_3", "Figure 4.3  Camera-based prototype verification page"),
    ("fig_4_4", "Figure 4.4  Event dashboard showing the active referendum"),
    ("fig_4_5", "Figure 4.5  Active referendum ballot page"),
    ("fig_4_6", "Figure 4.6  Vote review page"),
    ("fig_4_7", "Figure 4.7  Receipt page showing the Ballot ID and Proof Hash"),
    ("fig_4_8", "Figure 4.8  Public verification board"),
    ("fig_4_9", "Figure 4.9  Proof verification result"),
    ("fig_4_10", "Figure 4.10  Tally dashboard"),
    ("fig_4_11", "Figure 4.11  Admin login page"),
    ("fig_4_12", "Figure 4.12  Admin mock voter registry page"),
    ("fig_4_13", "Figure 4.13  Admin create-voter confirmation page"),
]

TABLE_ENTRIES = [
    ("tbl_1_1", "Table 1.1  Objectives-Methodology Mapping Table"),
    ("tbl_3_1", "Table 3.1  Functional Requirements of the System"),
    ("tbl_3_2", "Table 3.2  Non-Functional Requirements of the System"),
    ("tbl_3_3", "Table 3.3  Threat Model and Mitigation Strategy"),
    ("tbl_3_4", "Table 3.4  Voter Table Design"),
    ("tbl_3_5", "Table 3.5  Event Table Design"),
    ("tbl_3_6", "Table 3.6  Ballot Table Design"),
    ("tbl_3_7", "Table 3.7  Public and Private Verification Data"),
    ("tbl_4_1", "Table 4.1  Hardware Requirements"),
    ("tbl_4_2", "Table 4.2  Software Requirements"),
    ("tbl_4_3", "Table 4.3  Evaluation Criteria"),
    ("tbl_4_4", "Table 4.4  Functional Evaluation Results"),
    ("tbl_4_5", "Table 4.5  Direct Winterfell Proof-Engine Benchmark Summary"),
    ("tbl_4_6", "Table 4.6  Comparative Benchmark Against Existing Systems"),
    ("tbl_4_7", "Table 4.7  Program Interfaces"),
]

ABBREVIATIONS = [
    ("API", "Application Programming Interface"),
    ("BVAS", "Bimodal Voter Accreditation System"),
    ("CIS", "Computer and Information Sciences"),
    ("CSS", "Cascading Style Sheets"),
    ("FYP", "Final Year Project"),
    ("HTML", "HyperText Markup Language"),
    ("INEC", "Independent National Electoral Commission"),
    ("NIMC", "National Identity Management Commission"),
    ("NIN", "National Identification Number"),
    ("RAM", "Random Access Memory"),
    ("SQL", "Structured Query Language"),
    ("STARK", "Scalable Transparent Argument of Knowledge"),
    ("UI", "User Interface"),
    ("URL", "Uniform Resource Locator"),
    ("ZKP", "Zero-Knowledge Proof"),
]

REFERENCES = [
    "Adida, B. (2008). Helios: Web-based open-audit voting. In Proceedings of the 17th USENIX Security Symposium (pp. 335-348).",
    "Ali, S. T., & Murray, J. (2016). An overview of end-to-end verifiable voting systems. arXiv. https://arxiv.org/abs/1605.08554",
    "Alsadi, M., Casey, M., Dragan, C. C., Dupressoir, F., Riley, L., Sallal, M., Schneider, S., Treharne, H., Wadsworth, J., & Wright, P. (2019). Towards end-to-end verifiable online voting: Adding verifiability to established voting systems. arXiv. https://arxiv.org/abs/1912.00288",
    "Ben-Sasson, E., Bentov, I., Horesh, Y., & Riabzev, M. (2018). Scalable, transparent, and post-quantum secure computational integrity. IACR ePrint Archive, 2018/046. https://eprint.iacr.org/2018/046",
    "Chondros, N., Zhang, B., Zacharias, T., Diamantopoulos, P., Maneas, S., Patsonakis, C., Delis, A., Kiayias, A., & Roussopoulos, M. (2015). D-DEMOS: A distributed, end-to-end verifiable, internet voting system. arXiv. https://arxiv.org/abs/1507.06812",
    "Federal Republic of Nigeria. (2022). Electoral Act, 2022. Government Printer.",
    "Goldwasser, S., Micali, S., & Rackoff, C. (1989). The knowledge complexity of interactive proof systems. SIAM Journal on Computing, 18(1), 186-208.",
    "International IDEA. (2007). Voting from abroad: The International IDEA handbook. International Institute for Democracy and Electoral Assistance.",
    "McMurtry, E., Boyen, X., Culnane, C., Gjøsteen, K., Haines, T., & Teague, V. (2021). Towards verifiable remote voting with paper assurance. arXiv. https://arxiv.org/abs/2111.04210",
    "Quaglia, E. A., & Smyth, B. (2017). A short introduction to secrecy and verifiability for elections. arXiv. https://arxiv.org/abs/1702.03168",
    "The Belenios voting system. (n.d.). https://www.belenios.org/",
]

TABLE_BOOKMARKS = {
    "Table 1.1: Objectives-Methodology Mapping Table": "tbl_1_1",
    "Table 3.1: Functional Requirements of the System": "tbl_3_1",
    "Table 3.2: Non-Functional Requirements of the System": "tbl_3_2",
    "Table 3.3: Threat Model and Mitigation Strategy": "tbl_3_3",
    "Table 3.4: Voter Table Design": "tbl_3_4",
    "Table 3.5: Event Table Design": "tbl_3_5",
    "Table 3.6: Ballot Table Design": "tbl_3_6",
    "Table 3.7: Public and Private Verification Data": "tbl_3_7",
    "Table 4.1: Hardware Requirements": "tbl_4_1",
    "Table 4.2: Software Requirements": "tbl_4_2",
    "Table 4.3: Evaluation Criteria": "tbl_4_3",
    "Table 4.4: Functional Evaluation Results": "tbl_4_4",
    "Table 4.6: Comparative Benchmark Against Existing Systems": "tbl_4_6",
    "Table 4.7: Program Interfaces": "tbl_4_7",
}


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_backup() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    if OUTPUT_PATH.exists():
        backup = BACKUP_DIR / f"Defense Ready Final Year Project Report - pre-python-rebuild {OUTPUT_PATH.stat().st_mtime_ns}.docx"
        shutil.copy2(OUTPUT_PATH, backup)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        body.remove(child)
    body.append(sect_pr)


def add_bookmark(paragraph, name: str) -> None:
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "0")
    bookmark_start.set(qn("w:name"), name)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "0")
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_base_styles(doc: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(12)


def configure_section(section, start: int | None, fmt: str | None, show_page_number: bool) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.left_margin = Cm(3.0)

    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para._element.getparent().remove(para._element)

    if show_page_number:
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(para, "PAGE")

    sect_pr = section._sectPr
    for node in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(node)
    if start is not None or fmt is not None:
        pg = OxmlElement("w:pgNumType")
        if start is not None:
            pg.set(qn("w:start"), str(start))
        if fmt is not None:
            pg.set(qn("w:fmt"), fmt)
        sect_pr.append(pg)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    end_run = paragraph.add_run(" ")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def format_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5, before=0, after=6, first_line=0.5):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(first_line) if first_line else None


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    format_paragraph(p)
    return p


def add_center_paragraph(doc: Document, text: str, *, bold=False, spacing=1.0, after=0, before=0, size=12):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=spacing, before=before, after=after, first_line=0)
    return p


def add_heading(doc: Document, text: str, level: int, centered: bool = False):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if level > 1 else 14)
    format_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT,
        spacing=1.0,
        before=6,
        after=6,
        first_line=0,
    )
    return p


def add_table_caption(doc: Document, text: str, bookmark: str | None = None):
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, before=6, after=3, first_line=0)
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_figure_caption(doc: Document, text: str, bookmark: str | None = None):
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, before=3, after=9, first_line=0)
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_image(doc: Document, image_path: Path, caption: str, bookmark: str | None = None, width_inches: float = 6.0) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, before=3, after=3, first_line=0)
    add_figure_caption(doc, caption, bookmark=bookmark)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if i == 0:
                        run.bold = True
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_pageref_entry(doc: Document, text: str, bookmark: str) -> None:
    p = doc.add_paragraph(style="Normal")
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=3, first_line=0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.add_run(text)
    p.add_run("\t")
    add_field(p, f"PAGEREF {bookmark} \\h")


def normalize_text(text: str) -> str:
    replacements = {
        "PROTOYPE": "PROTOTYPE",
        "prototypeâ€™s": "prototype’s",
        "Table 4.5A: Direct Winterfell proof-engine benchmark summary": "Table 4.5: Direct Winterfell Proof-Engine Benchmark Summary",
        "Table 4.5: Comparative Benchmark Against Existing Systems": "Table 4.6: Comparative Benchmark Against Existing Systems",
        "Table 4.6: Program Interfaces": "Table 4.7: Program Interfaces",
        "Table 4.6.": "Table 4.7.",
        "Table 4.6 ": "Table 4.7 ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def add_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_center_paragraph(doc, "DESIGN AND IMPLEMENTATION OF A SECURE DIASPORA E-VOTING PROTOTYPE USING zk-STARK PROTOCOLS", bold=True, spacing=1.0, after=12)
    for _ in range(3):
        doc.add_paragraph()
    add_center_paragraph(doc, "BY", bold=True, spacing=1.0)
    for _ in range(3):
        doc.add_paragraph()
    add_center_paragraph(doc, "ADELUKA OLUWAMBEPELUMI EMMANUEL", bold=True, spacing=1.0)
    add_center_paragraph(doc, "Matric No.: ____________________", spacing=1.0, after=18)
    add_center_paragraph(doc, "A PROJECT SUBMITTED TO THE DEPARTMENT OF COMPUTER AND INFORMATION SCIENCES, COLLEGE OF SCIENCE AND TECHNOLOGY, COVENANT UNIVERSITY, OTA, OGUN STATE.", spacing=1.0, after=12)
    add_center_paragraph(doc, "IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF THE BACHELOR OF SCIENCE (HONOURS) DEGREE IN COMPUTER SCIENCE.", spacing=1.0, after=12)
    for _ in range(4):
        doc.add_paragraph()
    add_center_paragraph(doc, "JUNE, 2026", bold=True, spacing=1.0)


def add_preliminary_pages(doc: Document) -> None:
    add_heading(doc, "CERTIFICATION", 1, centered=True)
    add_body_paragraph(doc, "I hereby certify that this project was carried out by Adeluka Oluwambepelumi Emmanuel in the Department of Computer and Information Sciences, College of Science and Technology, Covenant University, Ota, Ogun State, Nigeria, under appropriate academic supervision.")
    doc.add_paragraph()
    add_body_paragraph(doc, "______________________________")
    add_body_paragraph(doc, "Project Supervisor")
    doc.add_paragraph()
    add_body_paragraph(doc, "______________________________")
    add_body_paragraph(doc, "Head of Department")

    doc.add_page_break()
    add_heading(doc, "DEDICATION", 1, centered=True)
    add_body_paragraph(doc, "This work is dedicated to God Almighty, whose grace sustained the entire study, and to my family for their encouragement, sacrifices, and unwavering support throughout the programme.")

    doc.add_page_break()
    add_heading(doc, "ACKNOWLEDGEMENTS", 1, centered=True)
    add_body_paragraph(doc, "My sincere gratitude goes to God for His grace, strength, and guidance throughout the course of this project. I am also grateful to my family and loved ones for their patience, encouragement, and support during the research, implementation, and documentation stages of this work.")
    add_body_paragraph(doc, "I further appreciate the academic guidance, technical feedback, and institutional support received from the Department of Computer and Information Sciences, Covenant University. The comments and reviews provided during the development of the DiasporaVote prototype contributed significantly to the clarity, scope control, and final presentation of this study.")

    doc.add_page_break()
    add_heading(doc, "TABLE OF CONTENT", 1, centered=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')

    doc.add_page_break()
    add_heading(doc, "LIST OF FIGURES", 1, centered=True)
    for bookmark, entry in FIGURE_ENTRIES:
        add_pageref_entry(doc, entry, bookmark)

    doc.add_page_break()
    add_heading(doc, "LIST OF TABLES", 1, centered=True)
    for bookmark, entry in TABLE_ENTRIES:
        add_pageref_entry(doc, entry, bookmark)

    doc.add_page_break()
    add_heading(doc, "ABBREVIATIONS", 1, centered=True)
    for abbr, meaning in ABBREVIATIONS:
        p = doc.add_paragraph(style="Normal")
        p.add_run(f"{abbr:<8}{meaning}")
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=3, first_line=0)

    doc.add_page_break()
    add_heading(doc, "ABSTRACT", 1, centered=True)
    p = doc.add_paragraph(style="Normal")
    p.add_run(
        "This study addressed the challenge of how secure diaspora voting could be demonstrated in a controlled Nigerian context without overclaiming institutional integration or national-election readiness. The project designed and implemented DiasporaVote, a secure binary referendum prototype that combined a React frontend, a Flask backend, SQLite persistence, a mock National Identification Number registry, and a Winterfell-based zk-STARK proof engine. The implemented workflow covered mock voter accreditation, camera-based prototype verification, event-aware ballot access, controlled Yes/No vote submission, proof generation and verification, encrypted ballot storage, receipt issuance, public-board publication of privacy-preserving verification metadata, and tally display for the active referendum event. System evaluation included registered-voter login, rejection of unregistered users, verification-state enforcement, valid and invalid ballot handling, duplicate-vote rejection, proof verification, tally behaviour, admin registry operations, and persistence checks. The results showed that the controlled binary referendum workflow was feasible within the scope of the prototype. The main contribution of the study was the implementation of a technically defensible diaspora voting prototype that demonstrated proof-backed ballot acceptance and server-mediated public verification while clearly stating its limitations, including mock identity assurance, binary-ballot scope, and non-production deployment status."
    )
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.0, after=6, first_line=0)
    keywords = doc.add_paragraph(style="Normal")
    run = keywords.add_run("Keywords: diaspora voting, e-voting, zk-STARK, zero-knowledge proof, public verification, referendum prototype")
    run.italic = True
    format_paragraph(keywords, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=0, first_line=0)


def add_benchmark_summary(doc: Document, benchmark_data: dict) -> None:
    rows = [[
        "Case",
        "Runs",
        "Avg. Gen. (ms)",
        "Avg. Ver. (ms)",
        "Avg. Proof Size (bytes)",
    ]]
    for item in benchmark_data["results"]:
        rows.append([
            item["synthetic_ballot_case"],
            str(item["run_count"]),
            f'{item["generation_time_ms"]["average"]:.4f}',
            f'{item["verification_time_ms"]["average"]:.4f}',
            str(item["proof_size_bytes"]["average"]),
        ])
    add_table_caption(doc, "Table 4.5: Direct Winterfell Proof-Engine Benchmark Summary", bookmark="tbl_4_5")
    add_table(doc, rows)


def add_references(doc: Document) -> None:
    add_heading(doc, "REFERENCES", 1, centered=True)
    for ref in REFERENCES:
        p = doc.add_paragraph(style="Normal")
        p.add_run(ref)
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=3, first_line=0)
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.63)


def add_chapter_five(doc: Document) -> None:
    add_heading(doc, "CHAPTER FIVE: SUMMARY, RECOMMENDATIONS, LIMITATIONS, AND CONCLUSION", 1, centered=True)
    add_heading(doc, "5.1 Summary", 2)
    add_body_paragraph(doc, "This study implemented a secure diaspora e-voting prototype for a controlled binary referendum using zk-STARK protocols. The completed system combined mock voter accreditation, camera-based prototype verification, event-aware ballot handling, proof-backed vote acceptance, encrypted storage, public receipt publication, server-mediated verification, and tally display within a single workflow designed for academic demonstration.")
    add_body_paragraph(doc, "The project also showed that the trust problem in remote voting is broader than interface design alone. Registry checks, duplicate-vote control, proof generation, proof hashing, privacy-preserving public verification metadata, and aggregate tally logic all had to operate together before the prototype could make a credible security claim.")
    add_body_paragraph(doc, "From a final year project perspective, the study further demonstrated that a tightly controlled problem definition was necessary to produce an academically defensible implementation. Limiting the ballot model to a binary referendum kept the proof logic, workflow checks, and evaluation evidence aligned with what could be implemented and defended honestly.")

    add_heading(doc, "5.2 Recommendations", 2)
    recommendations = [
        "Independent client-side or auditor-side proof verification should be added so that verification does not depend entirely on the election server.",
        "Stronger identity assurance should be explored only under proper institutional, legal, and privacy approval.",
        "The proof workflow and surrounding backend logic should undergo a formal cryptographic and security audit.",
        "Larger-scale performance testing should be carried out with higher ballot volumes and concurrent users.",
        "Future work should address coercion resistance, usability evaluation, and broader public-audit models.",
    ]
    for i, item in enumerate(recommendations, start=1):
        add_body_paragraph(doc, f"{i}. {item}")
    add_body_paragraph(doc, "These recommendations are intentionally sequenced around trustworthiness: the first priority is reducing dependence on server-side claims, followed by improving identity assurance, and then stress-testing the combined cryptographic and operational workflow.")

    add_heading(doc, "5.3 Limitations of the Study", 2)
    limitations = [
        "The identity layer used a mock NIN registry and did not connect to national identity infrastructure.",
        "The system did not integrate with official electoral infrastructure or any official national election backend.",
        "The camera-based verification step was a prototype presence check and not certified biometric identity matching.",
        "Public verification was server-mediated and did not yet provide fully independent local verification.",
        "The ballot model was restricted to a controlled binary referendum.",
        "The prototype was not production-ready for national elections.",
    ]
    for i, item in enumerate(limitations, start=1):
        add_body_paragraph(doc, f"{i}. {item}")

    add_heading(doc, "5.4 Conclusion", 2)
    add_body_paragraph(doc, "The study achieved its main aim of designing and implementing a secure diaspora e-voting prototype using zk-STARK protocols. Within the limits of a final year technical project, the implemented system demonstrated that proof-backed ballot acceptance, controlled accreditation, privacy-preserving public verification metadata, and event-aware tallying can be combined into a coherent remote-voting workflow. The strongest conclusion of the work is therefore not that national diaspora voting has been solved, but that a careful, scope-controlled prototype can make the security and verification problem more concrete, measurable, and academically defensible for future research and development.")
    add_body_paragraph(doc, "Accordingly, the final position of the report is that the prototype is valuable as a proof-of-concept and research instrument, not as a production election platform. Its contribution lies in showing a credible path for combining remote-voting workflow controls with zk-STARK-backed acceptance logic in a Nigerian diaspora-voting context.")


def add_body_content(doc: Document, ordered: dict, benchmark_data: dict) -> None:
    pending_table_caption = None
    pending_table_bookmark = None
    current_heading = ""
    objective_index = 0
    benchmark_inserted = False

    chapter_figure_inserts = {
        "3.2 Requirement Analysis": [
            ("Figure 3.1: Use-case diagram for DiasporaVote actors", DIAGRAMS_DIR / "use_case_diagram.png", "fig_3_1", "Figure 3.1 illustrates the main actors and interactions in the DiasporaVote workflow, including the voter, administrator, and public verification layer."),
        ],
        "3.3 System Architecture": [
            ("Figure 3.2: DiasporaVote system architecture", DIAGRAMS_DIR / "system_architecture.png", "fig_3_2", "Figure 3.2 shows the layered architecture that connects the React frontend, Flask backend, proof engine, storage layer, and public verification components."),
        ],
        "3.5 Voting Protocol Design": [
            ("Figure 3.3: Voting workflow for the controlled referendum prototype", DIAGRAMS_DIR / "voting_workflow.png", "fig_3_3", "Figure 3.3 traces the controlled referendum workflow from accreditation and verification through vote submission, proof generation, and tallying."),
        ],
        "3.6 zk-STARK Constraint System Design": [
            ("Figure 3.4: zk-STARK proof verification flow", DIAGRAMS_DIR / "zk_stark_verification_flow.png", "fig_3_4", "Figure 3.4 summarises how the proof artifact is generated, verified, hashed, and linked to public verification metadata."),
        ],
        "3.7 Database Design": [
            ("Figure 3.5: Database and entity-relationship design", DIAGRAMS_DIR / "database_er_diagram.png", "fig_3_5", "Figure 3.5 presents the database entities and relationships that support voters, events, ballots, and public verification data."),
        ],
    }

    section_46_figure_map = {
        "The system was divided into modules": [
            ("Figure 4.1: Landing page of the DiasporaVote prototype", SCREENSHOTS_DIR / "figure_4_1_landing_page.png", "fig_4_1", "Figure 4.1 shows the landing page that introduces the prototype and routes users into the accreditation and voting workflow."),
        ],
        "The voter accreditation module handled": [
            ("Figure 4.2: Mock NIN accreditation page", SCREENSHOTS_DIR / "figure_4_2_accreditation_page.png", "fig_4_2", "Figure 4.2 shows the accreditation interface through which a voter submits a mock NIN for registry validation."),
        ],
        "The camera-based verification module handled": [
            ("Figure 4.3: Camera-based prototype verification page", SCREENSHOTS_DIR / "figure_4_3_camera_verification_page.png", "fig_4_3", "Figure 4.3 shows the camera-based prototype verification stage that must be completed before ballot access is granted."),
        ],
        "The event module handled": [
            ("Figure 4.4: Event dashboard showing the active referendum", SCREENSHOTS_DIR / "figure_4_4_event_dashboard.png", "fig_4_4", "Figure 4.4 shows the event dashboard that exposes the active referendum context for the current voting session."),
        ],
        "The ballot module handled": [
            ("Figure 4.5: Active referendum ballot page", SCREENSHOTS_DIR / "figure_4_5_ballot_page.png", "fig_4_5", "Figure 4.5 shows the binary ballot interface for the active referendum event."),
            ("Figure 4.6: Vote review page", SCREENSHOTS_DIR / "figure_4_6_vote_review_page.png", "fig_4_6", "Figure 4.6 shows the vote review interface presented immediately before final ballot submission."),
        ],
        "The proof module handled": [
            ("Figure 4.7: Receipt page showing the Ballot ID and Proof Hash", SCREENSHOTS_DIR / "figure_4_7_receipt_page.png", "fig_4_7", "Figure 4.7 shows the receipt page that returns the Ballot ID and Proof Hash after successful proof-backed ballot acceptance."),
        ],
        "The public verification module handled": [
            ("Figure 4.8: Public verification board", SCREENSHOTS_DIR / "figure_4_8_public_verification_board.png", "fig_4_8", "Figure 4.8 shows the public verification board that exposes privacy-preserving receipt metadata."),
            ("Figure 4.9: Proof verification result", SCREENSHOTS_DIR / "figure_4_9_proof_verification_result.png", "fig_4_9", "Figure 4.9 shows the proof verification response returned when a stored proof is checked through the verification interface."),
        ],
        "The tally module handled": [
            ("Figure 4.10: Tally dashboard", SCREENSHOTS_DIR / "figure_4_10_tally_dashboard.png", "fig_4_10", "Figure 4.10 shows the tally dashboard that aggregates accepted Yes and No votes for the active referendum."),
        ],
        "The admin module handled": [
            ("Figure 4.11: Admin login page", SCREENSHOTS_DIR / "figure_4_11_admin_login.png", "fig_4_11", "Figure 4.11 shows the administrator login interface for protected registry management access."),
            ("Figure 4.12: Admin mock voter registry page", SCREENSHOTS_DIR / "figure_4_12_admin_registry.png", "fig_4_12", "Figure 4.12 shows the mock voter registry management page used to review and maintain eligible voter records."),
            ("Figure 4.13: Admin create-voter confirmation page", SCREENSHOTS_DIR / "figure_4_13_admin_create_voter.png", "fig_4_13", "Figure 4.13 shows the confirmation interface used after a new mock voter record is created."),
        ],
    }

    for element in ordered["elements"]:
        if element["type"] == "paragraph":
            text = normalize_text((element.get("text") or "").strip())
            if not text:
                continue

            if text in CHAPTER_MAP:
                if doc.paragraphs and doc.paragraphs[-1].text.strip():
                    doc.add_page_break()
                add_heading(doc, CHAPTER_MAP[text], 1, centered=True)
                current_heading = CHAPTER_MAP[text]
                objective_index = 0
                continue

            if text in SKIP_PARAGRAPHS:
                continue

            if re.fullmatch(r"\d+\.\d+(?:\.\d+)?\s+.+", text):
                level = 3 if text.count(".") >= 2 else 2
                add_heading(doc, text, level)
                current_heading = text
                objective_index = 0

                if text in chapter_figure_inserts:
                    for caption, image_path, bookmark, intro in chapter_figure_inserts[text]:
                        add_body_paragraph(doc, intro)
                        add_image(doc, image_path, caption, bookmark=bookmark)

                if text == "4.5.2 Evaluation Results" and not benchmark_inserted:
                    add_benchmark_summary(doc, benchmark_data)
                    benchmark_inserted = True
                continue

            if text.startswith("Table "):
                pending_table_caption = text
                pending_table_bookmark = TABLE_BOOKMARKS.get(text)
                continue

            if current_heading == "1.3 Aim and Objectives of the Study" and text in OBJECTIVE_TEXTS:
                objective_index += 1
                add_body_paragraph(doc, f"{objective_index}. {text}")
                continue

            add_body_paragraph(doc, text)

            if current_heading == "4.6 Program Modules and Interfaces":
                for prefix, figures in section_46_figure_map.items():
                    if text.startswith(prefix):
                        for caption, image_path, bookmark, intro in figures:
                            add_body_paragraph(doc, intro)
                            add_image(doc, image_path, caption, bookmark=bookmark)
                        break

        elif element["type"] == "table":
            if pending_table_caption:
                add_table_caption(doc, pending_table_caption, bookmark=pending_table_bookmark)
                pending_table_caption = None
                pending_table_bookmark = None
            add_table(doc, element.get("rows", []))


def main() -> None:
    ensure_backup()
    ordered = read_json(ORDERED_JSON_PATH)
    benchmark_data = read_json(BENCHMARK_JSON_PATH)

    shutil.copy2(DRAFT_PATH, OUTPUT_PATH)
    doc = Document(str(OUTPUT_PATH))

    clear_document_body(doc)
    set_update_fields_on_open(doc)
    set_base_styles(doc)

    configure_section(doc.sections[0], start=None, fmt=None, show_page_number=False)
    add_cover_page(doc)

    prelim = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(prelim, start=2, fmt="lowerRoman", show_page_number=True)
    add_preliminary_pages(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body, start=1, fmt="decimal", show_page_number=True)
    add_body_content(doc, ordered, benchmark_data)
    doc.add_page_break()
    add_chapter_five(doc)
    doc.add_page_break()
    add_references(doc)

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
