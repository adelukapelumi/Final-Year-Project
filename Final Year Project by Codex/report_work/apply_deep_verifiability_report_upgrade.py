from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


PROJECT_DIR = Path(__file__).resolve().parents[1]
DOC_PATH = PROJECT_DIR / "Defense Ready Final Year Project Report.docx"
BACKUP_DIR = PROJECT_DIR / "report_work" / "backups"


def find_paragraph_contains(doc: Document, snippet: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if snippet in paragraph.text:
            return paragraph
    raise RuntimeError(f"Could not find paragraph containing: {snippet}")


def replace_paragraph_text(doc: Document, snippet: str, new_text: str) -> Paragraph:
    paragraph = find_paragraph_contains(doc, snippet)
    paragraph.text = new_text
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    new_paragraph.add_run(text)
    return new_paragraph


def main() -> int:
    backup_dir = BACKUP_DIR
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup_path = backup_dir / (
        "Defense Ready Final Year Project Report - pre-deep-verifiability-upgrade "
        f"{datetime.now().strftime('%Y-%m-%d-%H%M%S')}.docx"
    )
    shutil.copy2(DOC_PATH, backup_path)

    doc = Document(DOC_PATH)

    subsection_style = find_paragraph_contains(doc, "3.2.1 Functional Requirements").style.name
    section_style = find_paragraph_contains(doc, "4.5.4 Security Analysis").style.name
    body_style = find_paragraph_contains(
        doc, "This chapter presents the analysis and design of the secure diaspora e-voting prototype"
    ).style.name

    replace_paragraph_text(
        doc,
        "Chondros et al.",
        "Chondros et al. (2015) proposed D-DEMOS as a distributed, end-to-end verifiable internet voting system with a replicated bulletin board, distributed vote collection, and trustees that support privacy and result production under a stronger formal trust model. This is important to the present study because it shows that public verifiability in remote voting is not only about proving ballot validity; it also depends on how the bulletin board, audit path, and trust assumptions are organised across multiple components.",
    )
    replace_paragraph_text(
        doc,
        "A key strength of D-DEMOS is that it treats remote voting as a complete system problem",
        "A key strength of D-DEMOS is its distributed bulletin board and its stronger formal security framing. DiasporaVote does not claim D-DEMOS-level distributed trust, fault tolerance, or trustee-based decentralisation. Instead, DiasporaVote contributes a smaller but stronger academic prototype in a different direction: a focused zk-STARK-backed referendum workflow that now adds event-scoped nullifiers, vote commitments, append-only hash-chain bulletin-board integrity, and independent verification bundles. The comparison therefore clarifies scope rather than implying equivalence: D-DEMOS represents a broader distributed internet-voting architecture, while DiasporaVote demonstrates how deeper proof-backed verifiability can be integrated into a controlled diaspora referendum prototype.",
    )

    replace_paragraph_text(
        doc,
        "The zk-STARK constraint system design explains the rules that guide proof-backed ballot acceptance in the prototype.",
        "The zk-STARK constraint system design explains the rules that guide proof-backed ballot acceptance in the upgraded prototype. In the earlier version, the proof mainly demonstrated the shallow binary ballot relation v(v - 1) = 0. The deep verifiability upgrade strengthens that statement by moving from a proof that only validates a public vote bit to a proof that links public verification metadata to private witness values while still hiding the voter’s actual choice. The design goal is therefore no longer only ballot-domain correctness, but privacy-preserving evidence that the published receipt metadata was derived correctly from the accepted private ballot state.",
    )
    replace_paragraph_text(doc, "The basic ballot domain constraint is that the vote must be one of the allowed binary values.", "3.6.1 Deep Verifiability Upgrade")
    replace_paragraph_text(
        doc,
        "The eligibility and duplicate-vote rules are enforced at the system level before proof-backed storage.",
        "The private witness now includes the binary vote value, a private voter secret derived inside the backend trust boundary, and a private ballot salt. The public inputs expose only the event identifier, the published nullifier, and the published vote commitment. This means the proof no longer reveals the raw vote value to the verifier, but it still forces the witness to satisfy a binary-ballot constraint.",
    )
    replace_paragraph_text(
        doc,
        "The proof artifact is linked to the public verification process through a proof hash.",
        "The nullifier is event-scoped. It is derived from the private voter secret and the event identifier, so the same voter produces the same nullifier for the same event and a different nullifier for a different event. This allows the backend to reject duplicate ballots through a published value that does not reveal the voter’s NIN, NIN hash, or private secret.",
    )
    replace_paragraph_text(
        doc,
        "The proof verification process is server-mediated in the current implementation.",
        "The vote commitment is derived from the private binary vote value and a private ballot salt. The commitment becomes part of the public board and the independent verification bundle, while the vote value and salt remain private. This gives the proof deeper cryptographic meaning because the verifier now checks that the published commitment is consistent with some valid binary vote, rather than simply checking that a public bit lies in {0,1}.",
    )
    replace_paragraph_text(doc, "The design of the constraint system can be summarised as follows:", "3.6.2 Upgraded Proof Statement")
    replace_paragraph_text(
        doc,
        "The submitted vote must belong to the allowed binary ballot domain.",
        "The upgraded proof statement can be summarised in three parts. First, the private vote value must belong to the binary domain {0,1}. Second, the published nullifier must be correctly derived from the private voter secret and the event identifier. Third, the published vote commitment must be correctly derived from the private vote value and the private ballot salt.",
    )
    replace_paragraph_text(
        doc,
        "The voter must pass mock registry accreditation before ballot access.",
        "A field-friendly prototype hash was used inside the Winterfell constraint system rather than SHA-256. SHA-256 remains outside the circuit for proof-artifact hashing and public-board chain hashing. This division is important because it keeps the STARK circuit practical while still ensuring that external publication and bundle verification use a conventional digest for file integrity and append-only board chaining.",
    )
    replace_paragraph_text(
        doc,
        "The voter session must pass camera-based prototype verification before vote submission.",
        "Eligibility, active-session checks, and event status are still enforced at the backend and database layers before proof-backed storage. The zk-STARK proof therefore strengthens recorded ballot meaning, while the surrounding system still handles registry control, session state, and event access as part of the wider trust architecture.",
    )
    replace_paragraph_text(
        doc,
        "The voter must not have previously voted in the same event.",
        "3.6.3 Public Verifiability Linkage",
    )
    replace_paragraph_text(
        doc,
        "The proof engine must generate a proof artifact for the accepted ballot workflow.",
        "Each accepted ballot is linked to a proof hash, a public nullifier, a public vote commitment, and an append-only board chain. The backend also exports an independent verification bundle that contains public inputs, proof bytes, proof hash, public receipt metadata, and chain metadata without exposing raw votes, voter identity, NIN values, salts, private secrets, or private proof paths.",
    )
    replace_paragraph_text(
        doc,
        "The backend must verify the proof before storing the encrypted ballot.",
        "An external verifier can therefore re-check the proof outside Flask, confirm that the proof hash matches the proof artifact, confirm that the public inputs match the receipt metadata, and recompute the chain hash for the published ballot record. This changes the role of the proof engine from decorative evidence to an independently checkable acceptance artifact.",
    )
    replace_paragraph_text(
        doc,
        "The stored ballot record must be linked to a public proof hash.",
        "3.6.4 Prototype Limitation",
    )
    replace_paragraph_text(
        doc,
        "The public board must expose only privacy-preserving receipt metadata.",
        "The in-circuit hash used in this version is a field-friendly prototype construction rather than a production-standard Poseidon or Rescue deployment. This limitation is stated openly in the report. However, the prototype still improves the technical depth of the project because the proof now binds public receipt fields to private witness values instead of exposing only a shallow public binary check.",
    )
    replace_paragraph_text(
        doc,
        "This design keeps the project honest and defensible.",
        "This design keeps the project honest and defensible. It does not claim that all election properties are solved entirely inside the STARK circuit, but it does prove a substantially richer statement than the earlier version and integrates that statement with duplicate-vote control, public-board integrity, and independent verification.",
    )

    replace_paragraph_text(
        doc,
        "The database design describes how the prototype stores the data required for mock voter management",
        "The database design describes how the prototype stores the data required for mock voter management, event management, ballot storage, proof-linked receipts, public verification, independent bundle export, and tallying. The deep verifiability upgrade extends the schema so that each accepted ballot now carries a published nullifier, a published vote commitment, a private ballot salt, verification status metadata, and append-only hash-chain values. SQLite remains appropriate for the prototype because it keeps the implementation lightweight while still allowing event-aware ballot isolation and deterministic board-chain recomputation.",
    )
    replace_paragraph_text(
        doc,
        "The voter table stores mock eligible voter records.",
        "The voter table now stores a voter-secret hash rather than exposing any private voter secret. The ballot table now stores nullifier, vote_commitment, ballot_salt, previous_chain_hash, current_record_hash, chain_hash, proof_hash, proof_path, and public_inputs. These additions are important because they separate private backend-only material from public verification material while still allowing independent auditors to check whether the published receipt metadata is internally consistent. The public board exposes only Ballot ID, event information, nullifier, vote commitment, proof hash, timestamp, verification status, and chain values; it does not expose raw vote values, decrypted votes, voter identity, NIN values, salts, private secrets, session tokens, or private proof paths.",
    )

    benchmark_anchor = find_paragraph_contains(
        doc,
        "This makes the prototype useful as an academic demonstration of a proof-backed remote voting workflow.",
    )
    if "4.5.3.1 Deep Evaluation Results" not in "\n".join(p.text for p in doc.paragraphs):
        benchmark_anchor = insert_paragraph_after(benchmark_anchor, "4.5.3.1 Deep Evaluation Results", subsection_style)
        benchmark_anchor = insert_paragraph_after(
            benchmark_anchor,
            "The deep verifiability upgrade was evaluated with four dedicated scripts and an independent CLI verifier. In the proof-engine volume benchmark, 1,000 proofs were completed successfully with an average proof-generation time of 0.698148 ms, an average proof-verification time of 0.446392 ms, an average proof size of about 6.31 KB, and a 100% success rate. Larger requested scales of 10,000, 100,000, and 1,000,000 proofs were configured in the script but were skipped during this workstation run when the projected runtime exceeded the safety window.",
            body_style,
        )
        benchmark_anchor = insert_paragraph_after(
            benchmark_anchor,
            "In the backend ballot acceptance benchmark, 1,000 end-to-end ballots were accepted successfully with zero failures, an average ballot acceptance time of 81.117444 ms, and a throughput of 12.324884 ballots per second. The final tally remained correct at 500 Yes and 500 No ballots, and the benchmark confirmed nullifier uniqueness correctness for all accepted ballots. Higher benchmark targets of 10,000 and 100,000 ballots were kept in the script but skipped during this run under the same runtime-safety policy.",
            body_style,
        )
        benchmark_anchor = insert_paragraph_after(
            benchmark_anchor,
            "In the concurrent voting stress test, the prototype completed the 10-voter concurrency level with 100% success, 100% duplicate-rejection correctness, correct final tallying, average end-to-end request latency of 1284.17472 ms, and throughput of 5.366095 ballots per second. Higher concurrency levels of 50, 100, 200, and 500 voters remain available in the script for future runs on stronger hardware or with a longer benchmark window.",
            body_style,
        )
        benchmark_anchor = insert_paragraph_after(
            benchmark_anchor,
            "In the public board and database scalability test, the prototype completed 1,000-ballot and 10,000-ballot database sizes using isolated databases. At 1,000 ballots, tally query time was 37.6041 ms and chain verification time was 142.2365 ms. At 10,000 ballots, tally query time was 290.9437 ms and chain verification time was 1164.8683 ms, while first-page and paginated board queries remained in the low-millisecond range. This shows that the append-only board chain introduces measurable but manageable verification cost within the prototype scope.",
            body_style,
        )
        benchmark_anchor = insert_paragraph_after(benchmark_anchor, "4.5.3.2 Independent CLI Verification Result", subsection_style)
        benchmark_anchor = insert_paragraph_after(
            benchmark_anchor,
            "A verification bundle exported from an accepted ballot was checked outside the Flask server using the standalone command python tools/verify_bundle.py benchmark_results/sample_verification_bundle.json. The CLI verifier reported passed results for proof verification, proof-hash validation, receipt-consistency validation, and chain-hash validation, with a final verdict of valid. This is a major upgrade over the earlier server-mediated-only design because it demonstrates that the public receipt can now be checked independently of the main voting API.",
            body_style,
        )

    replace_paragraph_text(
        doc,
        "The security analysis examined how the implemented system addressed the main risks identified in the threat model.",
        "The upgraded security analysis is structured around core verifiability properties rather than only around isolated backend controls. For each property, the report states whether it is achieved, partially achieved, or not achieved, together with the main evidence and the main limitation.",
    )
    replace_paragraph_text(
        doc,
        "The second security property was duplicate-vote prevention.",
        "1. Cast-as-intended: partially achieved. Evidence: the voter reviews the binary choice before submission, the backend rejects invalid vote values, and the proof forces the committed vote witness to be binary. Limitation: the prototype does not yet provide an independent voter-side interface for verifying commitment construction before submission, so cast-as-intended assurance still relies partly on the trusted application workflow.",
    )
    replace_paragraph_text(
        doc,
        "The third security property was ballot validity.",
        "2. Recorded-as-cast: achieved within prototype scope. Evidence: each accepted ballot produces a Ballot ID, nullifier, vote commitment, proof hash, public-board publication, and verification bundle; the CLI verifier can confirm proof validity, proof-hash integrity, receipt consistency, and chain-hash consistency outside Flask. Limitation: the trust model is still single-server and does not yet include a distributed bulletin board.",
    )
    replace_paragraph_text(
        doc,
        "The fourth security property was privacy-preserving public verification.",
        "3. Tallied-as-recorded: achieved within prototype scope. Evidence: tallying is computed from accepted encrypted ballot records, benchmark runs preserved tally correctness at 1,000 backend ballots and in the concurrent stress test, and chain verification confirms that the published board record is internally consistent. Limitation: the tally still depends on backend decryption and does not yet implement a trust-minimised cryptographic tally scheme.",
    )
    replace_paragraph_text(
        doc,
        "The fifth security property was tamper-evidence.",
        "4. Ballot privacy: partially achieved. Evidence: the public board exposes nullifier, vote commitment, proof hash, timestamp, and chain metadata without publishing raw vote values, decrypted votes, full NINs, NIN hashes, salts, private secrets, session tokens, or private proof paths. Limitation: the backend still stores encrypted ballots and performs decryption for tallying, so the design is privacy-preserving but not yet end-to-end secret-ballot perfect in the strongest formal sense.",
    )
    replace_paragraph_text(
        doc,
        "The sixth security property was tally correctness within the prototype scope.",
        "5. Duplicate-vote prevention: achieved within prototype scope. Evidence: accepted ballots now derive an event-scoped nullifier from a private voter secret and event identifier, the ballot table enforces nullifier uniqueness, duplicate submissions are rejected, and the backend benchmark confirmed nullifier uniqueness correctness. Limitation: the current model still depends on backend-controlled secret derivation rather than a fully decentralised credential system.",
    )
    replace_paragraph_text(
        doc,
        "The final security point concerns the limit of server-mediated verification.",
        "6. Public board integrity: achieved within prototype scope. Evidence: every accepted ballot stores previous_chain_hash, current_record_hash, and chain_hash derived from canonical public ballot records; the board verification endpoint recomputes the chain from a fixed genesis hash; the scalability benchmark showed successful chain verification at 1,000 and 10,000 ballots. Limitation: integrity is append-only within one prototype database and is not yet backed by distributed replication or trustee consensus.",
    )

    independent_anchor = find_paragraph_contains(
        doc,
        "6. Public board integrity: achieved within prototype scope.",
    )
    insert_paragraph_after(
        independent_anchor,
        "7. Independent verifiability: partially achieved. Evidence: the project now exports verification bundles and provides a standalone CLI verifier that runs outside Flask and validates the Winterfell proof, proof hash, receipt consistency, and chain hash. Limitation: a browser/WASM verifier is not yet deployed, and the broader trust model is still smaller than systems such as D-DEMOS that use distributed bulletin boards and stronger formal decentralisation assumptions.",
        body_style,
    )

    replace_paragraph_text(
        doc,
        "This study implemented a secure diaspora e-voting prototype for a controlled binary referendum using zk-STARK protocols.",
        "This study implemented a secure diaspora e-voting prototype for a controlled binary referendum using zk-STARK protocols and then upgraded it for deeper verifiability. The completed system now combines mock voter accreditation, camera-based prototype verification, event-aware ballot handling, proof-backed vote acceptance, encrypted storage, public receipt publication, nullifier-based duplicate-vote prevention, vote commitments, append-only board chaining, and independent verification-bundle export within a single academic workflow.",
    )
    replace_paragraph_text(
        doc,
        "The project also showed that the trust problem in remote voting is broader than interface design alone.",
        "The project also demonstrated that the trust problem in remote voting is broader than interface design alone. The upgraded prototype shows how backend controls, event checks, binary-vote proof constraints, nullifier derivation, vote commitments, proof hashes, public-board chaining, independent CLI verification, and tally logic can work together as a layered trust architecture. This makes the proof engine materially meaningful rather than merely decorative, while still keeping the scope limited to an academically honest referendum prototype.",
    )
    replace_paragraph_text(
        doc,
        "1. Independent client-side or auditor-side proof verification should be added so that verification does not depend entirely on the election server.",
        "1. The prototype field-friendly hash used inside the Winterfell constraint system should be upgraded to a more standard audited permutation such as Poseidon or Rescue in a future implementation.",
    )
    replace_paragraph_text(
        doc,
        "2. Stronger identity assurance should be explored only under proper institutional, legal, and privacy approval.",
        "2. Stronger identity assurance should be explored only under proper institutional, legal, and privacy approval, especially if future work moves beyond a mock NIN registry.",
    )
    replace_paragraph_text(
        doc,
        "3. The proof workflow and surrounding backend logic should undergo a formal cryptographic and security audit.",
        "3. The proof workflow, nullifier derivation, board-chain design, and surrounding backend logic should undergo a formal cryptographic and security audit.",
    )
    replace_paragraph_text(
        doc,
        "4. Larger-scale performance testing should be carried out with higher ballot volumes and concurrent users.",
        "4. Larger-scale performance testing should be carried out with higher ballot volumes, longer safe runtime windows, and stronger hardware so that the already prepared benchmark scripts can be executed at their larger targets.",
    )
    replace_paragraph_text(
        doc,
        "5. Future work should address coercion resistance, usability evaluation, and broader public-audit models.",
        "5. Additional research should address coercion resistance, usability evaluation, and a browser-based or WASM verifier so that independent verification becomes easier for ordinary voters.",
    )
    recommendation_five = find_paragraph_contains(
        doc,
        "5. Additional research should address coercion resistance, usability evaluation, and a browser-based or WASM verifier",
    )
    if "6. Deployment hardening and broader trust distribution should be improved" not in "\n".join(
        paragraph.text for paragraph in doc.paragraphs
    ):
        insert_paragraph_after(
            recommendation_five,
            "6. Deployment hardening and broader trust distribution should be improved before any real-world pilot use is considered, including stronger board replication and clearer external-auditor workflows.",
            body_style,
        )
    replace_paragraph_text(
        doc,
        "1. The identity layer used a mock NIN registry and did not connect to national identity infrastructure.",
        "1. The identity layer still uses a mock NIN registry and does not connect to national identity infrastructure.",
    )
    replace_paragraph_text(
        doc,
        "2. The system did not integrate with official electoral infrastructure or any official national election backend.",
        "2. The system still does not integrate with official electoral infrastructure or any official national election backend.",
    )
    replace_paragraph_text(
        doc,
        "3. The camera-based verification step was a prototype presence check and not certified biometric identity matching.",
        "3. The camera-based verification step remains a prototype presence check and not certified biometric identity matching.",
    )
    replace_paragraph_text(
        doc,
        "4. Public verification was server-mediated and did not yet provide fully independent local verification.",
        "4. Independent verification has improved through exported bundles and the standalone CLI verifier, but a browser/WASM verifier is not yet deployed for ordinary voter-side verification.",
    )
    replace_paragraph_text(
        doc,
        "5. The ballot model was restricted to a controlled binary referendum.",
        "5. The in-circuit field-friendly hash is a prototype construction and not yet a full production-standard audited voting hash design.",
    )
    replace_paragraph_text(
        doc,
        "6. The prototype was not production-ready for national elections.",
        "6. The ballot model is still restricted to a controlled binary referendum and does not yet address richer ballot formats.",
    )
    limitation_six = find_paragraph_contains(
        doc,
        "6. The ballot model is still restricted to a controlled binary referendum and does not yet address richer ballot formats.",
    )
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "7. The system does not implement national collation, multi-level aggregation" not in all_text:
        limitation_six = insert_paragraph_after(
            limitation_six,
            "7. The system does not implement national collation, multi-level aggregation, or polling-unit-to-state result flow, and it still relies on a single-server SQLite trust boundary rather than a distributed bulletin board or trustee network.",
            body_style,
        )
    if "8. The prototype is not production-ready for national elections." not in all_text:
        insert_paragraph_after(
            limitation_six,
            "8. The prototype is not production-ready for national elections.",
            body_style,
        )
    replace_paragraph_text(
        doc,
        "The study achieved its main aim of designing and implementing a secure diaspora e-voting prototype using zk-STARK protocols.",
        "The study achieved its main aim of designing and implementing a secure diaspora e-voting prototype using zk-STARK protocols. More importantly, the upgraded version shows that the proof engine can carry deeper cryptographic meaning by linking public nullifiers and public vote commitments to private witness values while preserving a controlled binary-vote privacy model. Within final year project limits, the prototype now demonstrates cast validation, duplicate-vote resistance, append-only board integrity, and independent verifier support in a technically stronger way than the earlier shallow statement.",
    )

    doc.save(DOC_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
